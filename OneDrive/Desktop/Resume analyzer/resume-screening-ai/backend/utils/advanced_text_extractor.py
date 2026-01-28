"""
Advanced Text Extraction Module
Uses PyMuPDF (better than PyPDF2) and python-docx with intelligent cleaning
"""

import fitz  # PyMuPDF
import docx
import re
from pathlib import Path
from typing import Optional, Tuple
import logging

logger = logging.getLogger(__name__)


class AdvancedTextExtractor:
    """
    Extract text from PDFs and DOCX files with advanced cleaning
    """
    
    @staticmethod
    def clean_text(text: str) -> str:
        """
        Clean extracted text:
        - Remove extra whitespaces
        - Normalize line breaks
        - Remove special characters (but keep important ones)
        - Fix common OCR errors
        """
        if not text:
            return ""
        
        # Remove null bytes
        text = text.replace('\x00', '')
        
        # Normalize whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Normalize line breaks (keep paragraph structure)
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        # Remove page numbers (standalone numbers)
        text = re.sub(r'^\s*\d+\s*$', '', text, flags=re.MULTILINE)
        
        # Remove footer/header artifacts (common patterns)
        text = re.sub(r'Page \d+ of \d+', '', text, flags=re.IGNORECASE)
        text = re.sub(r'Confidential|Private|Internal Use Only', '', text, flags=re.IGNORECASE)
        
        # Fix common OCR errors
        text = text.replace('|', 'I')  # Often | is read instead of I
        text = text.replace('0', 'O') if text.isupper() else text  # O/0 confusion in uppercase
        
        # Remove URLs (keep domain names)
        text = re.sub(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+', 
                     '', text)
        
        # Clean up
        text = text.strip()
        
        return text
    
    @staticmethod
    def extract_from_pdf(file_path: str) -> Tuple[str, bool]:
        """
        Extract text from PDF using PyMuPDF
        Returns: (text, success)
        """
        try:
            text_chunks = []
            
            # Open PDF
            doc = fitz.open(file_path)
            
            # Extract from each page
            for page_num, page in enumerate(doc):
                # Get text
                page_text = page.get_text("text")  # "text" mode is best for resumes
                
                if page_text.strip():
                    text_chunks.append(page_text)
                else:
                    # If no text, try blocks (better for formatted PDFs)
                    blocks = page.get_text("blocks")
                    page_text = "\n".join([block[4] for block in blocks if len(block) > 4])
                    text_chunks.append(page_text)
            
            doc.close()
            
            # Combine all pages
            full_text = "\n\n".join(text_chunks)
            
            # Clean text
            full_text = AdvancedTextExtractor.clean_text(full_text)
            
            if len(full_text) < 50:
                logger.warning(f"Extracted text too short ({len(full_text)} chars). Might be image-based PDF.")
                return full_text, False
            
            logger.info(f"Successfully extracted {len(full_text)} characters from PDF")
            return full_text, True
            
        except Exception as e:
            logger.error(f"Error extracting from PDF: {e}")
            return "", False
    
    @staticmethod
    def extract_from_docx(file_path: str) -> Tuple[str, bool]:
        """
        Extract text from DOCX
        Returns: (text, success)
        """
        try:
            doc = docx.Document(file_path)
            
            text_chunks = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_chunks.append(para.text.strip())
            
            # Extract from tables (resumes often have tables)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_chunks.append(cell.text.strip())
            
            # Combine
            full_text = "\n".join(text_chunks)
            
            # Clean
            full_text = AdvancedTextExtractor.clean_text(full_text)
            
            if len(full_text) < 50:
                logger.warning(f"Extracted text too short ({len(full_text)} chars)")
                return full_text, False
            
            logger.info(f"Successfully extracted {len(full_text)} characters from DOCX")
            return full_text, True
            
        except Exception as e:
            logger.error(f"Error extracting from DOCX: {e}")
            return "", False
    
    @staticmethod
    def extract_text(file_path: str) -> str:
        """
        Auto-detect file type and extract text
        """
        path = Path(file_path)
        extension = path.suffix.lower()
        
        if extension == '.pdf':
            text, success = AdvancedTextExtractor.extract_from_pdf(file_path)
        elif extension in ['.docx', '.doc']:
            text, success = AdvancedTextExtractor.extract_from_docx(file_path)
        else:
            logger.error(f"Unsupported file type: {extension}")
            return ""
        
        if not success:
            logger.error(f"Failed to extract text from {file_path}")
        
        return text
